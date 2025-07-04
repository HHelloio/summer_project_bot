import telebot
from telebot import types
import io
from PyPDF2 import PdfReader
import ollama
import time
import torch
import chardet
from docx import Document
import numpy as np
from sentence_transformers import SentenceTransformer
import faiss
import os
from tqdm import tqdm  # Для отображения прогресса
from dotenv import load_dotenv



load_dotenv('token.env')
token = os.getenv("TELEGRAM_BOT_TOKEN")
bot = telebot.TeleBot(token=token)

# Глобальные переменные
all_texts = []
saved_requests = []
text_chunks = []
faiss_index = None
model_name = "gemma3:12b"

# Инициализация модели для эмбеддингов
embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
EMBEDDING_DIM = 384

# ======== НАСТРОЙКИ РЕЖИМА ========
DEBUG_MODE = False  # Включаем для диагностики


# ==================================

def check_gpu_support():
    pytorch_cuda = torch.cuda.is_available()
    ollama_gpu = False
    try:
        model_info = ollama.show(model_name)
        ollama_gpu = "gpu" in model_info.get('details', {}).get('backend', 'cpu')
    except Exception as e:
        if DEBUG_MODE:
            print(f"Ошибка проверки Ollama GPU: {e}")
    return pytorch_cuda or ollama_gpu


gpu_available = check_gpu_support()
if DEBUG_MODE:
    print(f"GPU доступен: {gpu_available}")


def check_ollama_connection(retries=5, delay=3):
    for i in range(retries):
        try:
            ollama.list()
            if DEBUG_MODE:
                print("✅ Подключение к Ollama успешно установлено")
            return True
        except Exception as e:
            if DEBUG_MODE:
                print(f"⚠️ Попытка {i + 1}/{retries}: Ошибка подключения - {str(e)}")
            time.sleep(delay)
    if DEBUG_MODE:
        print("❌ Не удалось подключиться к Ollama")
    return False


if not check_ollama_connection() and DEBUG_MODE:
    print("Пожалуйста, убедитесь что сервер Ollama запущен (команда: ollama serve)")


def extract_text_from_txt(file_bytes: bytes) -> str:
    try:
        encoding_detection = chardet.detect(file_bytes)
        encoding = encoding_detection['encoding'] or 'utf-8'
        confidence = encoding_detection['confidence']

        if DEBUG_MODE:
            print(f"Определена кодировка: {encoding} (уверенность: {confidence:.2%})")

        text = file_bytes.decode(encoding)

        if '�' in text and confidence < 0.9:
            alternative_encodings = ['cp1251', 'iso-8859-5', 'koi8-r', 'cp866']
            for alt_enc in alternative_encodings:
                try:
                    text = file_bytes.decode(alt_enc)
                    if DEBUG_MODE:
                        print(f"Использована альтернативная кодировка: {alt_enc}")
                    break
                except:
                    continue
        return text

    except Exception as e:
        if DEBUG_MODE:
            print(f"Ошибка декодирования: {e}")
        try:
            return file_bytes.decode('utf-8', errors='ignore')
        except:
            return "Не удалось извлечь текст из файла"


def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        text = ""

        # Извлекаем основной текст
        for page in reader.pages:
            text += page.extract_text() or ""

            # Извлекаем аннотации (комментарии, заметки)
            if '/Annots' in page:
                for annot in page['/Annots']:
                    annot_obj = annot.get_object()
                    if '/Contents' in annot_obj:
                        text += "\n" + annot_obj['/Contents']

        return text

    except Exception as e:
        error_msg = f"Ошибка при обработке PDF: {str(e)}"
        if DEBUG_MODE:
            print(error_msg)
        return error_msg


def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        doc = Document(io.BytesIO(file_bytes))
        text = []

        # Основные параграфы
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)

        # Таблицы
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text)

        # Заголовки и колонтитулы
        for section in doc.sections:
            if section.header:
                for para in section.header.paragraphs:
                    if para.text.strip():
                        text.append(para.text)
            if section.footer:
                for para in section.footer.paragraphs:
                    if para.text.strip():
                        text.append(para.text)

        return "\n".join(text)

    except Exception as e:
        error_msg = f"Ошибка при обработке DOCX: {str(e)}"
        if DEBUG_MODE:
            print(error_msg)
        return error_msg


def save_request(message):
    global saved_requests
    saved_requests.append(message.text)
    if DEBUG_MODE:
        print(saved_requests)
    bot.send_message(message.chat.id, f"✅ Запрос сохранён:\n\n{message.text}")
    process_request(message)


def split_into_chunks(text: str, max_words: int = 400) -> list[str]:
    """Улучшенное разбиение с сохранением контекста"""
    words = text.split()
    chunks = []
    current_chunk = []
    current_length = 0

    for word in words:
        if current_length + 1 <= max_words:
            current_chunk.append(word)
            current_length += 1
        else:
            chunks.append(" ".join(current_chunk))
            current_chunk = [word]
            current_length = 1

    if current_chunk:
        chunks.append(" ".join(current_chunk))

    return chunks


def create_faiss_index(embeddings: np.ndarray):
    index = faiss.IndexFlatIP(EMBEDDING_DIM)
    faiss.normalize_L2(embeddings)
    index.add(embeddings)
    return index


def process_and_embed_chunks() -> bool:
    global all_texts, text_chunks, faiss_index

    if not all_texts:
        return False

    try:
        text_chunks = []
        for doc in all_texts:
            chunks = split_into_chunks(doc, max_words=400)
            text_chunks.extend(chunks)

        # Логирование для диагностики
        if DEBUG_MODE:
            total_words = sum(len(chunk.split()) for chunk in text_chunks)
            print(f"Всего слов в документах: {total_words}")
            print(f"Создано {len(text_chunks)} фрагментов")
            print(f"Размеры первых 5 фрагментов: {[len(chunk.split()) for chunk in text_chunks[:5]]}")
            if text_chunks:
                print(f"Первый фрагмент: {text_chunks[0][:200]}...")

        embeddings = embedding_model.encode(text_chunks)
        faiss_index = create_faiss_index(embeddings)

        if DEBUG_MODE:
            print(f"Индекс FAISS построен. Векторов: {faiss_index.ntotal}")
        return True

    except Exception as e:
        if DEBUG_MODE:
            print(f"Ошибка при обработке чанков: {e}")
        return False


def find_relevant_chunks(query: str, top_k: int = 5) -> list[str]:
    global faiss_index, text_chunks

    if faiss_index is None or len(text_chunks) == 0:
        return []

    query_embedding = embedding_model.encode([query])
    faiss.normalize_L2(query_embedding)

    # Увеличиваем top_k для большей полноты
    distances, indices = faiss_index.search(query_embedding, top_k * 2)

    # Фильтруем слишком нерелевантные фрагменты
    relevant_chunks = []
    for i, dist in zip(indices[0], distances[0]):
        if dist > 0.3:  # Порог сходства
            relevant_chunks.append(text_chunks[i])
            if len(relevant_chunks) >= top_k:
                break

    if DEBUG_MODE:
        print(f"Найдено {len(relevant_chunks)} релевантных фрагментов")
        print(f"Сходства: {distances[0][:len(relevant_chunks)]}")

    return relevant_chunks


def generate_response(prompt: str, context: str) -> str:
    try:
        full_prompt = (
            f"Ты — AI-ассистент. Ответь на вопрос используя контекст ниже.\n\n"
            f"Контекст:\n{context}\n\n"
            f"Вопрос: {prompt}\n\n"
            f"Ответ:"
        )

        options = {
            'temperature': 0.7,
            'num_predict': 3000,  # Увеличиваем лимит ответа
            'top_p': 0.9
        }

        if gpu_available:
            options['num_gpu'] = 50

        response = ollama.generate(
            model=model_name,
            prompt=full_prompt,
            options=options
        )

        answer = response['response'].strip()
        # Удаляем возможное дублирование промпта
        for phrase in [prompt, "Ответ:", "Вопрос:"]:
            if answer.startswith(phrase):
                answer = answer[len(phrase):].strip()

        return answer

    except Exception as e:
        if DEBUG_MODE:
            print(f"Ошибка генерации: {e}")
        return "⚠️ Произошла ошибка при генерации ответа. Попробуйте позже."


def process_request(message):
    global all_texts, saved_requests, faiss_index, text_chunks

    if not all_texts:
        bot.send_message(message.chat.id, "❌ Нет загруженных файлов для обработки запроса.")
        return

    if faiss_index is None or len(text_chunks) == 0:
        bot.send_message(message.chat.id,
                         "⏳ Создаю индекс для поиска... Это может занять время для больших документов.")
        if not process_and_embed_chunks():
            bot.send_message(message.chat.id, "❌ Ошибка при создании индекса поиска.")
            return

    prompt = saved_requests[-1]
    bot.send_message(message.chat.id, "🔍 Ищу релевантные фрагменты в документах...")
    relevant_chunks = find_relevant_chunks(prompt, top_k=5)  # Увеличиваем количество фрагментов

    if not relevant_chunks:
        bot.send_message(message.chat.id, "❌ Не найдено релевантных фрагментов для вашего запроса.")
        return

    context = "\n\n".join(relevant_chunks)

    if DEBUG_MODE:
        mode = "GPU" if gpu_available else "CPU"
        bot.send_message(message.chat.id, f"⚙️ Обрабатываю запрос (на {mode})...")
        bot.send_message(message.chat.id, f"📚 Использую {len(relevant_chunks)} релевантных фрагментов")
    else:
        bot.send_message(message.chat.id, "⚙️ Обрабатываю запрос...")

    bot.send_message(message.chat.id, "🧠 Генерирую ответ...")
    response = generate_response(prompt, context)

    # Разбиваем длинный ответ на части
    max_length = 4000
    if len(response) > max_length:
        for i in range(0, len(response), max_length):
            bot.send_message(message.chat.id, response[i:i + max_length])
    else:
        bot.send_message(message.chat.id, f"📝 Ответ на ваш запрос:\n\n{response}")


@bot.message_handler(commands=['start'])
def start(message):
    markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
    btn1 = types.KeyboardButton("👋 И тебе не хворать")
    markup.add(btn1)
    bot.send_message(message.chat.id,
                     "👋 Привет! Я готов к работе. Отправь мне текстовые документы (PDF, DOCX, TXT), а затем задавай вопросы по их содержимому.",
                     reply_markup=markup)


@bot.message_handler(content_types=['text', 'document'])
def handle_files(message):
    global all_texts, text_chunks, faiss_index, saved_requests

    if message.document:
        file_info = bot.get_file(message.document.file_id)
        file_name = message.document.file_name
        file_bytes = bot.download_file(file_info.file_path)

        try:
            if file_name.lower().endswith('.txt'):
                text = extract_text_from_txt(file_bytes)
            elif file_name.lower().endswith('.pdf'):
                text = extract_text_from_pdf(file_bytes)
            elif file_name.lower().endswith(('.docx', '.doc')):
                text = extract_text_from_docx(file_bytes)
            else:
                bot.reply_to(message, '❌ Поддерживаются только файлы .txt, .pdf и .docx')
                return

            all_texts.append(text)
            text_chunks = []
            faiss_index = None

            # Логирование информации о файле
            if DEBUG_MODE:
                print(f"Загружено документов: {len(all_texts)}")
                print(f"Размер текста: {len(text.split())} слов")
                print(f"Первые 200 символов: {text[:200]}...")

            bot.reply_to(message, f"📄 Файл успешно загружен: {file_name}\n\nТекст извлечен: {len(text.split())} слов")

        except Exception as er:
            error_msg = f"⚠️ Ошибка при обработке файла: {er}"
            if DEBUG_MODE:
                error_msg += "\n\nПодсказка: Файл может иметь нестандартную кодировку или быть поврежден"
            bot.reply_to(message, error_msg)

    elif message.text:
        if message.text == '👋 И тебе не хворать':
            markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
            btn1 = types.KeyboardButton("Добавить файл")
            btn2 = types.KeyboardButton('Загрузить файл')
            btn3 = types.KeyboardButton('Удалить все файлы')
            btn4 = types.KeyboardButton("📝 Написать запрос")
            btn6 = types.KeyboardButton("💾 Сохранить индекс")
            btn7 = types.KeyboardButton("📥 Загрузить индекс")

            if DEBUG_MODE:
                btn5 = types.KeyboardButton("ℹ️ Информация о системе")
                markup.add(btn1, btn2, btn3, btn4, btn5, btn6, btn7)
            else:
                markup.add(btn1, btn2, btn3, btn4)

            bot.send_message(message.chat.id, 'Выберите действие:', reply_markup=markup)

        elif message.text == 'Добавить файл':
            bot.send_message(message.chat.id, 'Присылайте файл в формате PDF, DOCX или TXT')

        elif message.text == 'Загрузить файл':
            if not all_texts:
                bot.send_message(message.chat.id, "❌ Нет файлов для построения индекса.")
            else:
                bot.send_message(message.chat.id,
                                 "⏳ Создаю индекс для поиска... Это может занять время для больших документов.")
                if process_and_embed_chunks():
                    bot.send_message(message.chat.id, f"✅ Индекс создан! Обработано {len(text_chunks)} фрагментов")
                else:
                    bot.send_message(message.chat.id, "❌ Ошибка при создании индекса")

        elif message.text == 'Удалить все файлы':
            all_texts = []
            saved_requests = []
            text_chunks = []
            faiss_index = None
            bot.send_message(message.chat.id, '✅ Все файлы, запросы и индексы удалены')
            
            #Для удаления сохраненных индексов на диске
            try:
                os.remove("faiss_index.index")
            except FileNotFoundError:
                pass

            try:
                os.remove("text_chunks.pkl")
            except FileNotFoundError:
                pass

        elif message.text == '📝 Написать запрос':
            bot.send_message(message.chat.id, "Напиши свой запрос по содержимому документов:")
            bot.register_next_step_handler(message, save_request)

        elif message.text == '💾 Сохранить индекс':
            if faiss_index is None:
                bot.send_message(message.chat.id, "❌ Индекс не создан")
            else:
                try:
                    faiss.write_index(faiss_index, "faiss_index.index")
                    bot.send_message(message.chat.id, "✅ Индекс сохранен в файл faiss_index.index")
                except Exception as e:
                    bot.send_message(message.chat.id, f"❌ Ошибка сохранения: {str(e)}")

        elif message.text == '📥 Загрузить индекс':
            if not os.path.exists("faiss_index.index"):
                bot.send_message(message.chat.id, "❌ Файл индекса не найден")
            else:
                try:
                    # Используем глобальную переменную, объявленную в начале функции
                    faiss_index = faiss.read_index("faiss_index.index")

                    # Восстанавливаем text_chunks из файла (для примера - в реальности нужно сохранять отдельно)
                    bot.send_message(message.chat.id,
                                     "✅ Индекс загружен из файла\n⚠️ Текстовые фрагменты не восстановлены, нужно перестроить индекс")
                except Exception as e:
                    bot.send_message(message.chat.id, f"❌ Ошибка загрузки: {str(e)}")

        elif message.text == 'ℹ️ Информация о системе' and DEBUG_MODE:
            try:
                pytorch_info = f"PyTorch: {torch.__version__}\nCUDA: {torch.cuda.is_available()}"
                emb_info = f"Embedding model: {embedding_model._modules['0'].auto_model.config._name_or_path}"
                faiss_info = f"FAISS: {faiss_index.ntotal} векторов" if faiss_index else "FAISS: индекс не создан"
                model_info = ollama.show(model_name)
                parameters = model_info.get('parameters', 'неизвестно')

                info_msg = (
                    "⚙️ Информация о системе:\n"
                    f"• Модель: {model_name}\n"
                    f"• Режим: {'GPU' if gpu_available else 'CPU'}\n"
                    f"• Параметры: {parameters}\n"
                    f"• {emb_info}\n"
                    f"• {faiss_info}\n\n"
                    f"PyTorch информация:\n{pytorch_info}"
                )

                bot.send_message(message.chat.id, info_msg)
            except Exception as e:
                bot.send_message(message.chat.id, f"⚠️ Ошибка получения информации: {str(e)}")


bot.polling(none_stop=True, interval=0)