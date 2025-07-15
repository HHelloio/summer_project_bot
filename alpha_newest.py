import telebot
from telebot import types
import io
from PyPDF2 import PdfReader
import ollama
import time
import torch
import chardet
from docx import Document
import numpy as np
from sentence_transformers import SentenceTransformer
import faiss
import os
from dotenv import load_dotenv
from duckduckgo_search import DDGS
import nltk
from nltk.tokenize import sent_tokenize
from sentence_transformers import util
import networkx as nx 
from dataclasses import dataclass
from typing import List


nltk.download('punkt')

# Загрузка переменных окружения
load_dotenv('token.env')
token = os.getenv("TELEGRAM_BOT_TOKEN")
bot = telebot.TeleBot(token=token)
#TELEGRAM_BOT_TOKEN=8068131419:AAF1kLt_l9GDOZrc4MRMMxGECLd_pXMzblQ

# Глобальные переменные
all_texts = []
saved_requests = []
text_chunks = []
faiss_index = None
model_name = "gemma3:12b"
WEB_SEARCH_ENABLED = True  # Флаг для включения/выключения веб-поиска
chunk_graph = None  # Глобальная переменная для хранения графа

@dataclass
class ChunkMeta:
    id: int
    doc_id: int
    text: str
    position: int  # Позиция чанка в документе

# Инициализация модели для эмбеддингов
embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
EMBEDDING_DIM = 384

# ======== НАСТРОЙКИ РЕЖИМА ========
DEBUG_MODE = True  # Включаем для диагностики


# ==================================

def check_gpu_support():
    pytorch_cuda = torch.cuda.is_available()
    ollama_gpu = False
    try:
        model_info = ollama.show(model_name)
        ollama_gpu = "gpu" in model_info.get('details', {}).get('backend', 'cpu')
    except Exception as e:
        if DEBUG_MODE:
            print(f"Ошибка проверки Ollama GPU: {e}")
    return pytorch_cuda or ollama_gpu

gpu_available = check_gpu_support()
if DEBUG_MODE:
    print(f"GPU доступен: {gpu_available}")

def check_ollama_connection(retries=5, delay=3):
    for i in range(retries):
        try:
            ollama.list()
            if DEBUG_MODE:
                print("✅ Подключение к Ollama успешно установлено")
            return True
        except Exception as e:
            if DEBUG_MODE:
                print(f"⚠️ Попытка {i + 1}/{retries}: Ошибка подключения - {str(e)}")
            time.sleep(delay)
    if DEBUG_MODE:
        print("❌ Не удалось подключиться к Ollama")
    return False

if not check_ollama_connection() and DEBUG_MODE:
    print("Пожалуйста, убедитесь что сервер Ollama запущен (команда: ollama serve)")

def extract_text_from_txt(file_bytes: bytes) -> str:
    try:
        encoding_detection = chardet.detect(file_bytes)
        encoding = encoding_detection['encoding'] or 'utf-8'
        confidence = encoding_detection['confidence']

        if DEBUG_MODE:
            print(f"Определена кодировка: {encoding} (уверенность: {confidence:.2%})")

        text = file_bytes.decode(encoding)

        if '�' in text and confidence < 0.9:
            alternative_encodings = ['cp1251', 'iso-8859-5', 'koi8-r', 'cp866']
            for alt_enc in alternative_encodings:
                try:
                    text = file_bytes.decode(alt_enc)
                    if DEBUG_MODE:
                        print(f"Использована альтернативная кодировка: {alt_enc}")
                    break
                except:
                    continue
        return text

    except Exception as e:
        if DEBUG_MODE:
            print(f"Ошибка декодирования: {e}")
        try:
            return file_bytes.decode('utf-8', errors='ignore')
        except:
            return "Не удалось извлечь текст из файла"


def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        text = ""

        # Извлекаем основной текст
        for page in reader.pages:
            text += page.extract_text() or ""

            # Извлекаем аннотации (комментарии, заметки)
            if '/Annots' in page:
                for annot in page['/Annots']:
                    annot_obj = annot.get_object()
                    if '/Contents' in annot_obj:
                        text += "\n" + annot_obj['/Contents']

        return text

    except Exception as e:
        error_msg = f"Ошибка при обработке PDF: {str(e)}"
        if DEBUG_MODE:
            print(error_msg)
        return error_msg


def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        doc = Document(io.BytesIO(file_bytes))
        text = []

        # Основные параграфы
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)

        # Таблицы
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text)

        # Заголовки и колонтитулы
        for section in doc.sections:
            if section.header:
                for para in section.header.paragraphs:
                    if para.text.strip():
                        text.append(para.text)
            if section.footer:
                for para in section.footer.paragraphs:
                    if para.text.strip():
                        text.append(para.text)

        return "\n".join(text)

    except Exception as e:
        error_msg = f"Ошибка при обработке DOCX: {str(e)}"
        if DEBUG_MODE:
            print(error_msg)
        return error_msg

def save_request(message):
    global saved_requests
    saved_requests.append(message.text)
    if DEBUG_MODE:
        print(saved_requests)
    bot.send_message(message.chat.id, f"✅ Запрос сохранён:\n\n{message.text}")
    process_request(message)

def split_into_chunks_sent_overlap(text: str, max_sentences: int = 10, overlap_sentences: int = 2) -> list[str]:
    sentences = sent_tokenize(text)
    chunks = []
    i = 0

    while i < len(sentences):
        end = i + max_sentences
        chunk = " ".join(sentences[i:end])
        chunks.append(chunk)
        i += max_sentences - overlap_sentences

    if DEBUG_MODE:
        print(f"[CHUNKING-SENT] Сгенерировано {len(chunks)} чанков с overlap={overlap_sentences}")
    return chunks

def split_into_chunks_word_overlap(text: str, max_words: int = 400, overlap_words: int = 50) -> list[str]:
    words = text.split()
    chunks = []
    i = 0

    while i < len(words):
        chunk = words[i:i + max_words]
        chunks.append(" ".join(chunk))
        i += max_words - overlap_words

    if DEBUG_MODE:
        print(f"[CHUNKING-WORD] Сгенерировано {len(chunks)} чанков с overlap={overlap_words} слов")
    return chunks

CHUNK_METHOD = "sentence"  # или "word"

def split_into_chunks(text: str, doc_id: int = 0, **kwargs) -> List[ChunkMeta]:
    global CHUNK_METHOD
    raw_chunks = []

    if CHUNK_METHOD == "sentence":
        raw_chunks = split_into_chunks_sent_overlap(text, max_sentences=10, overlap_sentences=2)
    elif CHUNK_METHOD == "word":
        raw_chunks = split_into_chunks_word_overlap(text, max_words=400, overlap_words=50)
    else:
        raise ValueError("Неподдерживаемый метод чанкинга")

    chunks_with_meta = [
        ChunkMeta(
            id=i,
            doc_id=doc_id,
            text=chunk,
            position=i
        )
        for i, chunk in enumerate(raw_chunks)
    ]

    if DEBUG_MODE:
        print(f"[CHUNK_META] Документ {doc_id}: создано {len(chunks_with_meta)} чанков с метаданными")

    return chunks_with_meta

def create_faiss_index(embeddings: np.ndarray):
    index = faiss.IndexFlatIP(EMBEDDING_DIM)
    faiss.normalize_L2(embeddings)
    index.add(embeddings)
    return index

def build_chunk_graph(chunks: list[ChunkMeta], threshold=0.7):
    G = nx.Graph()
    texts = [chunk.text for chunk in chunks]
    embeddings = embedding_model.encode(texts, convert_to_tensor=True)

    for i, chunk in enumerate(chunks):
        G.add_node(i, text=chunk.text, doc_id=chunk.doc_id, position=chunk.position)

    for i in range(len(chunks)):
        for j in range(i + 1, len(chunks)):
            sim = util.cos_sim(embeddings[i], embeddings[j]).item()
            if sim >= threshold:
                G.add_edge(i, j, weight=round(sim, 3))

    if DEBUG_MODE:
        print(f"[GRAPH] Построен граф с {len(G.nodes)} узлами и {len(G.edges)} рёбрами")
    return G

def process_and_embed_chunks() -> bool:
    global all_texts, text_chunks, faiss_index, chunk_graph

    if not all_texts:
        return False

    try:
        text_chunks = []
        for doc_id, doc_text in enumerate(all_texts):
            chunks = split_into_chunks(doc_text, doc_id=doc_id)
            text_chunks.extend(chunks)

        texts = [chunk.text for chunk in text_chunks]
        embeddings = embedding_model.encode(texts)

        faiss_index = create_faiss_index(embeddings)
        chunk_graph = build_chunk_graph(text_chunks)

        if DEBUG_MODE:
            print(f"Индекс FAISS построен. Векторов: {faiss_index.ntotal}")
        return True

    except Exception as e:
        if DEBUG_MODE:
            print(f"Ошибка при обработке чанков: {e}")
        return False

def find_relevant_chunks(query: str, top_k: int = 20, expand_with_graph: bool = True) -> list[str]:
    global faiss_index, text_chunks, chunk_graph

    if faiss_index is None or len(text_chunks) == 0:
        return []

    query_embedding = embedding_model.encode([query])
    faiss.normalize_L2(query_embedding)

    distances, indices = faiss_index.search(query_embedding, top_k * 2)

    selected_indices = []
    for i, dist in zip(indices[0], distances[0]):
        if dist > 0.1:
            selected_indices.append(i)
            if len(selected_indices) >= top_k:
                break

    if DEBUG_MODE:
        print(f"[FAISS] Найдено {len(selected_indices)} релевантных чанков")

    expanded_indices = set(selected_indices)

    # Расширение через граф
    if expand_with_graph and chunk_graph:
        for idx in selected_indices:
            neighbors = list(chunk_graph.neighbors(idx))
            for neighbor in neighbors:
                if neighbor not in expanded_indices:
                    expanded_indices.add(neighbor)
                    if DEBUG_MODE:
                        sim = chunk_graph[idx][neighbor]['weight']
                        print(f"[GRAPH] Добавлен сосед {neighbor} к чанку {idx} (сходство={sim:.3f})")

    # Сортировка по оригинальному индексу (можно заменить на другую стратегию)
    final_chunks = [text_chunks[i].text for i in sorted(expanded_indices)]

    if DEBUG_MODE:
        print(f"[CHUNKS] Финальное количество чанков после расширения: {len(final_chunks)}")

    return final_chunks

def perform_web_search(query: str, max_results: int = 3) -> str:
    """Выполняет поиск в интернете и возвращает форматированные результаты"""
    try:
        results = []
        with DDGS() as ddgs:
            for result in ddgs.text(query, max_results=max_results):
                results.append(f"• [{result['title']}]({result['href']})\n{result['body']}")

        return "\n\n".join(results) if results else "Информация по запросу в интернете не найдена."

    except Exception as e:
        if DEBUG_MODE:
            print(f"Ошибка веб-поиска: {e}")
        return ""

def generate_response(prompt: str, context: str, web_context: str = "") -> str:
    try:
        # Формируем полный промпт с двумя контекстами
        full_prompt = (
            f"Ты — AI-ассистент. Ответь на вопрос используя контекст ниже. Перепроверь корректность информации.\n\n"
            f"КОНТЕКСТ ИЗ ДОКУМЕНТОВ:\n{context}\n\n"
            f"КОНТЕКСТ ИЗ ИНТЕРНЕТА:\n{web_context}\n\n"
            f"Вопрос: {prompt}\n\n"
            f"Ответ:"
        )

        options = {
            'temperature': 0.7,
            'num_predict': 3000,  # Увеличиваем лимит ответа
            'top_p': 0.9
        }

        if gpu_available:
            options['num_gpu'] = 50

        response = ollama.generate(
            model=model_name,
            prompt=full_prompt,
            options=options
        )

        answer = response['response'].strip()
        # Удаляем возможное дублирование промпта
        for phrase in [prompt, "Ответ:", "Вопрос:"]:
            if answer.startswith(phrase):
                answer = answer[len(phrase):].strip()

        return answer

    except Exception as e:
        if DEBUG_MODE:
            print(f"Ошибка генерации: {e}")
        return "⚠️ Произошла ошибка при генерации ответа. Попробуйте позже."

def process_request(message):
    global all_texts, saved_requests, faiss_index, text_chunks, WEB_SEARCH_ENABLED

    if not all_texts:
        bot.send_message(message.chat.id, "❌ Нет загруженных файлов для обработки запроса.")
        return

    if faiss_index is None or len(text_chunks) == 0:
        bot.send_message(message.chat.id,
                         "⏳ Создаю индекс для поиска... Это может занять время для больших документов.")
        if not process_and_embed_chunks():
            bot.send_message(message.chat.id, "❌ Ошибка при создании индекса поиска.")
            return

    prompt = saved_requests[-1]
    bot.send_message(message.chat.id, "🔍 Ищу релевантные фрагменты в документах...")
    relevant_chunks = find_relevant_chunks(prompt, top_k=5)  # Увеличиваем количество фрагментов

    if not relevant_chunks:
        bot.send_message(message.chat.id, "❌ Не найдено релевантных фрагментов для вашего запроса.")
        context = ""
    else:
        context = "\n\n".join(relevant_chunks)

    # Веб-поиск
    web_context = ""
    if WEB_SEARCH_ENABLED:
        #bot.send_message(message.chat.id, "🌐 Ищу информацию в интернете...")
        web_context = perform_web_search(prompt, max_results=3)
        if DEBUG_MODE:
            print(f"Веб-контекст: {web_context[:500]}...")

    if DEBUG_MODE:
        mode = "GPU" if gpu_available else "CPU"
        bot.send_message(message.chat.id, f"⚙️ Обрабатываю запрос (на {mode})...")
        if context:
            bot.send_message(message.chat.id, f"📚 Использую {len(relevant_chunks)} релевантных фрагментов")
        if web_context:
            bot.send_message(message.chat.id, "🌐 Использую результаты веб-поиска")
    else:
        bot.send_message(message.chat.id, "⚙️ Обрабатываю запрос...")

    bot.send_message(message.chat.id, "🧠 Генерирую ответ...")
    response = generate_response(prompt, context, web_context)

    # Разбиваем длинный ответ на части
    max_length = 4000
    if len(response) > max_length:
        for i in range(0, len(response), max_length):
            bot.send_message(message.chat.id, response[i:i + max_length])
    else:
        bot.send_message(message.chat.id, f"📝 Ответ на ваш запрос:\n\n{response}")

def create_main_keyboard():
    markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
    btn1 = types.KeyboardButton("Добавить файл")
    btn2 = types.KeyboardButton('Загрузить файл')
    btn3 = types.KeyboardButton('Удалить все файлы')
    btn4 = types.KeyboardButton("📝 Написать запрос")
    btn6 = types.KeyboardButton("💾 Сохранить индекс")
    btn7 = types.KeyboardButton("📥 Загрузить индекс")
    # Кнопка для управления веб-поиском
    btn_web = types.KeyboardButton("🌐 Веб-поиск: Вкл" if WEB_SEARCH_ENABLED else "🌐 Веб-поиск: Выкл")

    if DEBUG_MODE:
        btn5 = types.KeyboardButton("ℹ️ Информация о системе")
        markup.add(btn1, btn2, btn3, btn4, btn5, btn6, btn7)
    else:
        markup.add(btn1, btn2, btn3)

    return markup

@bot.message_handler(commands=['start'])
def start(message):
    markup = create_main_keyboard()
    bot.send_message(message.chat.id,
                     "👋 Привет! Я готов к работе. Отправь мне текстовые документы (PDF, DOCX, TXT), а затем задавай вопросы по их содержимому.",
                     reply_markup=markup)

@bot.message_handler(content_types=['text', 'document'])
def handle_files(message):
    global all_texts, text_chunks, faiss_index, saved_requests, WEB_SEARCH_ENABLED

    if message.document:
        file_info = bot.get_file(message.document.file_id)
        file_name = message.document.file_name
        file_bytes = bot.download_file(file_info.file_path)

        try:
            if file_name.lower().endswith('.txt'):
                text = extract_text_from_txt(file_bytes)
            elif file_name.lower().endswith('.pdf'):
                text = extract_text_from_pdf(file_bytes)
            elif file_name.lower().endswith(('.docx', '.doc')):
                text = extract_text_from_docx(file_bytes)
            else:
                bot.reply_to(message, '❌ Поддерживаются только файлы .txt, .pdf и .docx')
                return

            all_texts.append(text)
            text_chunks = []
            faiss_index = None

            # Логирование информации о файле
            if DEBUG_MODE:
                print(f"Загружено документов: {len(all_texts)}")
                print(f"Размер текста: {len(text.split())} слов")
                print(f"Первые 200 символов: {text[:200]}...")

            bot.reply_to(message, f"📄 Файл успешно загружен: {file_name}\n\nТекст извлечен: {len(text.split())} слов")

        except Exception as er:
            error_msg = f"⚠️ Ошибка при обработке файла: {er}"
            if DEBUG_MODE:
                error_msg += "\n\nПодсказка: Файл может иметь нестандартную кодировку или быть поврежден"
            bot.reply_to(message, error_msg)

    elif message.text:
        if message.text == '👋 И тебе не хворать':
            markup = create_main_keyboard()
            bot.send_message(message.chat.id, 'Выберите действие:', reply_markup=markup)

        elif message.text == 'Добавить файл':
            bot.send_message(message.chat.id, 'Присылайте файл в формате PDF, DOCX или TXT')

        elif message.text == 'Загрузить файл':
            if not all_texts:
                bot.send_message(message.chat.id, "❌ Нет файлов для построения индекса.")
            else:
                bot.send_message(message.chat.id,
                                 "⏳ Создаю индекс для поиска... Это может занять время для больших документов.")
                if process_and_embed_chunks():
                    bot.send_message(message.chat.id, f"✅ Индекс создан! Обработано {len(text_chunks)} фрагментов")
                else:
                    bot.send_message(message.chat.id, "❌ Ошибка при создании индекса")

        elif message.text == 'Удалить все файлы':
            all_texts = []
            saved_requests = []
            text_chunks = []
            faiss_index = None
            bot.send_message(message.chat.id, '✅ Все файлы, запросы и индексы удалены')

            # Для удаления сохраненных индексов на диске
            try:
                os.remove("faiss_index.index")
            except FileNotFoundError:
                pass

            try:
                os.remove("text_chunks.pkl")
            except FileNotFoundError:
                pass

        elif message.text == '📝 Написать запрос':
            bot.send_message(message.chat.id, "Напиши свой запрос по содержимому документов:")
            bot.register_next_step_handler(message, save_request)
            


        elif message.text == '💾 Сохранить индекс':
            if faiss_index is None:
                bot.send_message(message.chat.id, "❌ Индекс не создан")
            else:
                try:
                    faiss.write_index(faiss_index, "faiss_index.index")
                    bot.send_message(message.chat.id, "✅ Индекс сохранен в файл faiss_index.index")
                except Exception as e:
                    bot.send_message(message.chat.id, f"❌ Ошибка сохранения: {str(e)}")

        elif message.text == '📥 Загрузить индекс':
            if not os.path.exists("faiss_index.index"):
                bot.send_message(message.chat.id, "❌ Файл индекса не найден")
            else:
                try:
                    # Используем глобальную переменную, объявленную в начале функции
                    faiss_index = faiss.read_index("faiss_index.index")
                    bot.send_message(message.chat.id,
                                     "✅ Индекс загружен из файла\n⚠️ Текстовые фрагменты не восстановлены, нужно перестроить индекс")
                except Exception as e:
                    bot.send_message(message.chat.id, f"❌ Ошибка загрузки: {str(e)}")

        elif message.text == 'ℹ️ Информация о системе' and DEBUG_MODE:
            try:
                pytorch_info = f"PyTorch: {torch.__version__}\nCUDA: {torch.cuda.is_available()}"
                emb_info = f"Embedding model: {embedding_model._modules['0'].auto_model.config._name_or_path}"
                faiss_info = f"FAISS: {faiss_index.ntotal} векторов" if faiss_index else "FAISS: индекс не создан"
                model_info = ollama.show(model_name)
                parameters = model_info.get('parameters', 'неизвестно')
                graph_info = f"Граф: {chunk_graph.number_of_nodes()} узлов, {chunk_graph.number_of_edges()} рёбер" if chunk_graph else "Граф: не построен"


                info_msg = (
                    "⚙️ Информация о системе:\n"
                    f"• Модель: {model_name}\n"
                    f"• Режим: {'GPU' if gpu_available else 'CPU'}\n"
                    f"• Параметры: {parameters}\n"
                    f"• {emb_info}\n"
                    f"• {faiss_info}\n\n"
                    f"PyTorch информация:\n{pytorch_info}"
                    f"• {graph_info}\n"
                )

                bot.send_message(message.chat.id, info_msg)
            except Exception as e:
                bot.send_message(message.chat.id, f"⚠️ Ошибка получения информации: {str(e)}")

        # Обработка кнопки веб-поиска
        elif message.text.startswith('🌐 Веб-поиск:'):
            WEB_SEARCH_ENABLED = not WEB_SEARCH_ENABLED
            status = "ВКЛЮЧЕН" if WEB_SEARCH_ENABLED else "ВЫКЛЮЧЕН"
            bot.send_message(message.chat.id, f"Веб-поиск: {status}", reply_markup=create_main_keyboard())


bot.polling(none_stop=True, interval=0) 