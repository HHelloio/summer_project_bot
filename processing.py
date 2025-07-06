import io
from PyPDF2 import PdfReader
from docx import Document
import chardet
from sentence_transformers import SentenceTransformer
import numpy as np
import faiss
import ollama
import torch
import os

embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
EMBEDDING_DIM = 384
model_name = "gemma3:12b"
DEBUG_MODE = False

def check_gpu_support():
    pytorch_cuda = torch.cuda.is_available()
    ollama_gpu = False
    try:
        model_info = ollama.show(model_name)
        ollama_gpu = "gpu" in model_info.get('details', {}).get('backend', 'cpu')
    except Exception as e:
        if DEBUG_MODE:
            print(f"Ошибка проверки Ollama GPU: {e}")
    return pytorch_cuda or ollama_gpu

def extract_text_from_txt(file_bytes: bytes) -> str:
    encoding_detection = chardet.detect(file_bytes)
    encoding = encoding_detection['encoding'] or 'utf-8'
    try:
        return file_bytes.decode(encoding)
    except:
        return file_bytes.decode('utf-8', errors='ignore')

def extract_text_from_pdf(file_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(file_bytes))
    return "\n".join([page.extract_text() or "" for page in reader.pages])

def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)

def split_into_chunks(text: str, max_words: int = 400) -> list[str]:
    words = text.split()
    return [" ".join(words[i:i+max_words]) for i in range(0, len(words), max_words)]

def create_faiss_index(embeddings: np.ndarray):
    faiss.normalize_L2(embeddings)
    index = faiss.IndexFlatIP(EMBEDDING_DIM)
    index.add(embeddings)
    return index

def embed_chunks(chunks: list[str]):
    return embedding_model.encode(chunks)

def find_relevant_chunks(query: str, index, chunks, top_k: int = 5):
    query_embedding = embedding_model.encode([query])
    faiss.normalize_L2(query_embedding)
    distances, indices = index.search(query_embedding, top_k * 2)
    return [chunks[i] for i, d in zip(indices[0], distances[0]) if d > 0.3][:top_k]

def generate_response(prompt: str, context: str, gpu_available: bool) -> str:
    full_prompt = f"Ты — AI-ассистент. Ответь на вопрос используя контекст ниже.\n\nКонтекст:\n{context}\n\nВопрос: {prompt}\n\nОтвет:"
    options = {'temperature': 0.7, 'num_predict': 3000, 'top_p': 0.9}
    if gpu_available:
        options['num_gpu'] = 50
    response = ollama.generate(model=model_name, prompt=full_prompt, options=options)
    answer = response['response'].strip()
    for phrase in [prompt, "Ответ:", "Вопрос:"]:
        if answer.startswith(phrase):
            answer = answer[len(phrase):].strip()
    return answer

def save_faiss_index(index, path="faiss_index.index"):
    faiss.write_index(index, path)

def load_faiss_index(path="faiss_index.index"):
    return faiss.read_index(path)